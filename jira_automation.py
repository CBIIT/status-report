import os
import requests
import json
from docx import Document
from dotenv import load_dotenv
from typing import List, Dict, Any

# Load environment variables
load_dotenv()

# Configuration
JIRA_TOKEN = os.getenv("JIRA_TOKEN")
JIRA_URL = os.getenv("JIRA_URL")
JQL = os.getenv("JIRA_JQL")

OLLAMA_URL = "http://localhost:11434/api/generate"


class JiraToDocxAutomation:
    """Main class for JIRA to DOCX automation system"""
    
    def __init__(self, project_name: str):
        self.validate_config()
        self.project_name = project_name
    
    def validate_config(self):
        """Validate that all required environment variables are set"""
        if not JIRA_TOKEN or JIRA_TOKEN == "your_token_here":
            raise ValueError("JIRA_TOKEN must be set in .env file")
        if not JIRA_URL or "yourdomain" in JIRA_URL:
            raise ValueError("JIRA_URL must be set to your actual JIRA instance in .env file")
        if not JQL:
            raise ValueError("JIRA_JQL must be set in .env file")

    def fetch_issues(self, project_name: str) -> List[Dict[str, Any]]:
        """
        Fetch issues from JIRA using JQL
        
        Returns:
            List of issue dictionaries
        """
        try:
            headers = {
                "Authorization": f"Bearer {JIRA_TOKEN}",
                "Accept": "application/json",
            }
            
            params = {
                "jql": "project = '" + project_name + "' " + JQL,
                "fields": "issuetype,key,summary,status,project,priority,assignee,reporter,description"
            }
            
            print(f"Fetching issues from JIRA with JQL: {JQL}")
            response = requests.get(
                "https://tracker.nci.nih.gov/rest/api/2/search", 
                headers=headers, 
                params=params,
                timeout=30
            )
            
            if response.status_code != 200:
                raise Exception(f"JIRA API error: {response.status_code} - {response.text}")
            
            data = response.json()
            issues = data.get("issues", [])
            print(f"Successfully fetched {len(issues)} issues from JIRA")
            return issues
            
        except requests.exceptions.RequestException as e:
            print(f"Error fetching issues from JIRA: {e}")
            return []
        except Exception as e:
            print(f"Unexpected error: {e}")
            return []
    
    def summarize_with_ollama(self, text: str) -> str:
        """
        Summarize text using Ollama LLM
        
        Args:
            text: The text to summarize
            
        Returns:
            Summarized text
        """
        try:
            # Prepare the request body for Ollama chat API
            body = {
                "model": "llama3",
                "prompt": (
                    "You are a project manager assistant. Given a list of JIRA issues or tasks with the fields: "
                    "Issue Type, Issue Key, Summary, and Status, create a concise and professional high level summary of planned "
                    "or ongoing activities for the current or upcoming month. Do not list individual issues. Give a brief overview."
                    f"Here is the list of issues: {text}"
                ),
                "stream": False
            }

            headers = {
                "Content-Type": "application/json"
            }
            
            print("Generating summary with Ollama...")
            response = requests.post(
                OLLAMA_URL, 
                json=body,
                headers=headers
            )
            
            if response.status_code != 200:
                print(f"Ollama API error: {response.status_code} - {response.text}")
                return f"Error generating summary: {response.text}"
            
            result = response.json()
            summary = result.get("response", {})
            return summary.strip()
            
        except requests.exceptions.RequestException as e:
            print(f"Error connecting to Ollama: {e}")
            return f"Error connecting to Ollama: {e}"
        except Exception as e:
            print(f"Unexpected error during summarization: {e}")
            return f"Error generating summary: {e}"


    def generate_word_document(self, issues, ai_summary, project_name, filename: str = "JIRA_Summary_Report.docx"):
        """
        Generate Word document with issues grouped by project
        
        Args:
            projects_data: Dictionary with project names as keys and project data as values
            filename: Output filename
        """
        try:
            print(f"Generating Word document: {filename}")
            
            doc = Document()
            
            # Add title
            title = doc.add_heading(f"{project_name}: Tasks completed or to be continued in the upcoming month.", 1)
            title.alignment = 1  # Center alignment
            
  
            # Process each project
            project_summary = ai_summary
                
            # Create table with headers
            issues_table = doc.add_table(rows=1, cols=4)
            issues_table.style = 'Table Grid'
                
            # Set table headers
            hdr_cells = issues_table.rows[0].cells
            hdr_cells[0].text = 'Issue Type'
            hdr_cells[1].text = 'Issue key'
            hdr_cells[2].text = 'Summary'
            hdr_cells[3].text = 'Status'
                
             # Add issues to table
            for issue in issues:
                    row_cells = issues_table.add_row().cells
                    row_cells[0].text = issue.get("issue type", "N/A")
                    row_cells[1].text = issue.get("issue key", "No summary")
                    row_cells[2].text = issue.get("summary", "No summary")
                    row_cells[3].text = issue.get("status", "No status")


                # Project summary section
            doc.add_heading("Project Summary", level=2)
            doc.add_paragraph(project_summary)
            doc.add_page_break()
            # Save document
            doc.save(filename)
            print(f"Document saved successfully as {filename}")
            
        except Exception as e:
            print(f"Error generating Word document: {e}")
    

    def run(self):
        """
        Main execution method
        """
        print("Starting JIRA to DOCX automation...")
        print("=" * 50)
        
        # Step 1: Fetch issues from JIRA
        issues = self.fetch_issues(self.project_name)
        if not issues:
            print("No issues found or error fetching issues. Exiting.")
            return
        
        # Step 2: Process each issue
        issue_summaries = []
        
        for i, issue in enumerate(issues, 1):
            print(f"\nProcessing issue {i}/{len(issues)}: {issue.get('key', 'Unknown')}")
            # Extract metadata
            # "fields": "issuetype,key,summary,status,project,priority,assignee,reporter,description"
            fields = issue.get("fields", {})
            issue_data = {
                "issue type": fields.get("issuetype", "Unknown").get("name", "Unknown"),
                "issue key": issue.get("key", "No key"),
                "summary": fields.get("summary", "No summary"),
                "status": fields.get("status", {}).get("name", "Unknown"),
            }
            
            issue_summaries.append(issue_data)
        
        # Step 3: Generate Word document
        print(f"\n{'='*50}")
        print(f"\n{'='*50}")
       
        issues_string = "\n".join([f"Issue Key: {issue['issue key']}, Summary: {issue['summary']}, Status: {issue['status']}" for issue in issue_summaries])
        print(issues_string)
        ai_summary = self.summarize_with_ollama(issues_string)
        print(f"\nAI Summary:\n{ai_summary}")

        self.generate_word_document(issue_summaries, ai_summary, self.project_name)

        # # Step 4: Group issues by project and generate project-level summaries
        # print(f"\n{'='*50}")
        # projects = self.group_issues_by_project(issues)
        # project_summaries = []
        
        # for project_name, project_issues in projects.items():
        #     # Summarize issues at the project level
        #     summary = self.summarize_project_issues(project_name, project_issues)
        #     project_summaries.append({
        #         "project_name": project_name,
        #         "summary": summary
        #     })
        
        # # Generate Word document for project summaries
        # project_summary_filename = "JIRA_Project_Summary_Report.docx"
        # self.generate_project_summary_document(project_summaries, project_summary_filename)
        
        print(f"\n{'='*50}")
        print("Automation completed successfully!")
        print(f"Generated report for {len(issue_summaries)} issues and project summaries.")
    
 
def main():
    """Entry point for the application"""
    try:
        automation = JiraToDocxAutomation("Index of NCI Studies")
        automation.run()
    except Exception as e:
        print(f"Error: {e}")
        print("\nPlease ensure:")
        print("1. Your .env file is properly configured")
        print("2. Ollama is running locally with llama3 model")
        print("3. Your JIRA credentials are correct")


if __name__ == "__main__":
    main()
